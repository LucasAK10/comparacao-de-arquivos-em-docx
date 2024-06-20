# Biblioteca padrão do Python para interações com o sistema operacional (Manipulação de arquivos e diretórios)
import os 

# Biblioteca padrão do Python para interações relacionadas ao tempo (medir intervalos de tempo)
import time

# Importa componentes essenciais do Flask, um micro-framework para criar aplicações web
from flask import Flask, request, render_template, send_from_directory, flash, redirect, url_for

# Função para garantir que os nomes dos arquivos enviados são seguros
from werkzeug.utils import secure_filename

# Biblioteca para agendamento de tarefas em segundo plano
from apscheduler.schedulers.background import BackgroundScheduler

# Importa a classe 'Document' da biblioteca 'python-docx' para manipulação
from docx import Document

# Importa a classe 'RGBcolor' para definir cores de texto nos documentos DOCX
from docx.shared import RGBColor

# Configurações do Aplicativo

# Define o diretório onde os arquivos serão armazenados
UPLOAD_FOLDER = 'uploads'

# Define as extensões de arquivos permitidas para o upload
ALLOWED_EXTENSIONS = {'docx'}

# Define o tamanho máximo permitido para os arquivos
MAX_CONTENT_LENGTH = 16 * 1024 * 1024  # 16 MB

# Define o tempo de vida útil dos arquivos no servidor
FILE_LIFETIME = 60 * 60 * 24  # 24 horas

# Iniciação do Flask

# Cria uma instância do Flask (Flask = biblioteca que permite criar um micro framework para criar aplicações web)
app = Flask(__name__)

# Configura o diretório de upload no Flask
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER

# Configura o tamanho máximo dos arquivos no Flask
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Define uma chave secreta para o Flask (usada para gerenciar sessões e cookies de forma segura)
app.secret_key = 'supersecretkey'  # Alterar para uma chave secreta adequada

# Função para verificar Extensões Permitas

# Verifica se o arquivo possui uma extensão permitida
def allowed_file(filename):

    # "'.' in filename": verifica se há um ponto no nome do arquivo
    # "file name.rsplit('.',1)[1].lower() in ALLOWED_EXTENSIONS": Divide o nome do arquivo pelo ponto, pega a útlima parte (extensão) e verifica se está nas extensões permitidas
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

# Função para ler o conteúdo de arquivo DOCX

# Lê conteúdo de um arquivo DOCX e retorna como texto
def read_docx(file_path):

# Abre o arquivo DOCX
    doc = Document(file_path)
# Inicializa uma lista para armazenar o texto dos parágrafos
    full_text = []
# Itera sobre os parágrafos do documentos
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def compare_documents(file1_path, file2_path, output_file):
    doc1 = Document(file1_path)
    doc2 = Document(file2_path)
    output_doc = Document()

    def mark_difference(text, color):
        run = output_doc.add_paragraph().add_run(text)
        run.font.color.rgb = RGBColor(*color)

    def compare_paragraphs(paras1, paras2):
        for p1, p2 in zip(paras1, paras2):
            if p1.text != p2.text:
                mark_difference(f'Document 1: {p1.text}', (255, 0, 0))
                mark_difference(f'Document 2: {p2.text}', (0, 255, 0))
            else:
                output_doc.add_paragraph(p1.text)

    compare_paragraphs(doc1.paragraphs, doc2.paragraphs)

    output_doc.save(output_file)

def clean_upload_folder():
    now = time.time()
    for filename in os.listdir(UPLOAD_FOLDER):
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        if os.path.isfile(file_path):
            file_age = now - os.path.getmtime(file_path)
            if file_age > FILE_LIFETIME:
                os.remove(file_path)
                print(f'Removido: {file_path}')

@app.route('/', methods=['GET', 'POST'])
def upload_files():
    if request.method == 'POST':
        if 'file1' not in request.files or 'file2' not in request.files:
            flash('Nenhum arquivo foi enviado.')
            return render_template('upload.html')
        file1 = request.files['file1']
        file2 = request.files['file2']

        if file1.filename == '' or file2.filename == '':
            flash('Nenhum arquivo foi selecionado.')
            return render_template('upload.html')

        if file1 and allowed_file(file1.filename) and file2 and allowed_file(file2.filename):
            filename1 = secure_filename(file1.filename)
            filename2 = secure_filename(file2.filename)
            file1_path = os.path.join(app.config['UPLOAD_FOLDER'], filename1)
            file2_path = os.path.join(app.config['UPLOAD_FOLDER'], filename2)
            output_file = os.path.join(app.config['UPLOAD_FOLDER'], 'relatorio_diferencas.docx')

            file1.save(file1_path)
            file2.save(file2_path)
            
            compare_documents(file1_path, file2_path, output_file)
            
            return redirect(url_for('result'))
        else:
            flash('Apenas arquivos DOCX são permitidos.')
            return render_template('upload.html')

    return render_template('upload.html')

@app.route('/result', methods=['GET', 'POST'])
def result():
    if request.method == 'POST':
        if 'new_upload' in request.form:
            return redirect(url_for('upload_files'))
        elif 'shutdown' in request.form:
            shutdown_server()
            return 'Servidor encerrado.'
    return render_template('result.html')

@app.route('/uploads/<filename>')
def uploaded_file(filename):
    return send_from_directory(app.config['UPLOAD_FOLDER'], filename)

def shutdown_server():
    func = request.environ.get('werkzeug.server.shutdown')
    if func is None:
        raise RuntimeError('Não é possível encerrar o servidor.')
    func()

if __name__ == '__main__':
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)

    scheduler = BackgroundScheduler()
    scheduler.add_job(clean_upload_folder, 'interval', hours=24)
    scheduler.start()

    try:
        app.run(debug=True)
    except (KeyboardInterrupt, SystemExit):
        scheduler.shutdown()


